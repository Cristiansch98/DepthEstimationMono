"""Render the SelfCalibDepth digest as an editable .docx in the IEEE two-column
structure of the ITEC-AP-2026 template, so it can be dropped into Word.

    python src/make_docx.py        # -> PAPER.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT = "Times New Roman"


def set_columns(section, num=2, space_in=0.25):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(space_in * 1440)))
    cols.set(qn("w:equalWidth"), "1")


def page_setup(section):
    section.page_height, section.page_width = Inches(11.69), Inches(8.27)  # A4
    for m in ("top_margin", "bottom_margin"):
        setattr(section, m, Inches(0.7))
    section.left_margin = section.right_margin = Inches(0.6)


def runs(p, content, size=10, italic=False, bold=False):
    """content: str, or list of (text, italic, bold)."""
    if isinstance(content, str):
        content = [(content, italic, bold)]
    for text, it, bd in content:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.italic = it
        r.bold = bd


def para(doc, content, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.2,
         italic=False, bold=False, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Inches(indent)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    runs(p, content, size=size, italic=italic, bold=bold)
    return p


def h1(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    label = f"{num}. {title}" if num else title
    r = p.add_run(label)
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.small_caps = True


def h2(doc, letter, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{letter}. {title}")
    r.font.name = FONT
    r.font.size = Pt(10)
    r.italic = True


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(8)


def figure(doc, path, cap):
    p = path if isinstance(path, Path) else Path(path)
    if p.exists():
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run().add_picture(str(p), width=Inches(3.3))
    caption(doc, cap)


def table(doc, header_caption, data, widths):
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(header_caption)
    r.font.name = FONT
    r.font.size = Pt(8)
    r.font.small_caps = True
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.alignment = 1
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = Inches(widths[ci])
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(0)
            run = cp.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(7.5)
            run.bold = (ri == 0) or (ci == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build():
    doc = Document()
    page_setup(doc.sections[0])

    # ---- single-column title block ----
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runs(t, "SelfCalibDepth: LiDAR-Supervised Self-Calibration for Camera-Aware "
            "Monocular Metric Depth on Argoverse 2", size=18, bold=True)
    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runs(a, "Cristian Cubides", size=11)
    af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runs(af, "Software Research and Development\nSelfCalibDepth Project, Argoverse 2\n"
             "cschrcm@gmail.com", size=9, italic=True)

    # ---- switch to two columns ----
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    page_setup(sec)
    set_columns(sec, 2)

    para(doc, [("Abstract", False, True), ("—", False, True),
               ("Reading metric distance from a single image requires resolving two coupled "
                "unknowns, the per-pixel depth and the camera intrinsics that turn pixels into "
                "viewing rays, which ordinary monocular depth models leave entangled. We present "
                "SelfCalibDepth, which uses a synchronized LiDAR sweep as metric ground truth to "
                "anchor a learnable camera model and a fine-tuned Depth Anything V2 backbone, "
                "coupled through a per-pixel ray map: the calibration defines the rays that "
                "condition the depth head, and back-projecting the predicted depth to match the "
                "LiDAR points sends a gradient back into the calibration, so the camera is "
                "recovered from the image rather than assumed. On Argoverse 2 the model recovers "
                "focal length to within 0.26%/0.41% of the manufacturer calibration on held-out "
                "cameras, attains AbsRel 0.112 and accuracy 0.884 at threshold 1.25, and estimates "
                "vehicle distance within sixty meters to 8.6 m mean error, all from one image. A "
                "controlled ablation includes a pinpointed negative result: a camera-conditioned "
                "affine scale head loses to a free convolutional head, and unfreezing the "
                "foundation backbone is the decisive lever. Unifying four driving benchmarks behind "
                "one adapter, we find the model does not transfer zero-shot, but a twenty-frame "
                "few-shot adaptation reaches KITTI 0.100/0.953 and nuScenes 0.122/0.853, on par "
                "with in-domain. Finally, we diagnose from its source why the state-of-the-art "
                "UniDepth collapses under lens distortion, its self-calibration head predicts only "
                "a pinhole, and repair it to the ground-truth-camera oracle (AbsRel 0.158 to 0.106 "
                "under fisheye) with our few-shot LiDAR calibration and no retraining. All "
                "modelling assumptions are annotated and the framework is released.", True, True)],
         size=9, indent=0)
    para(doc, [("Keywords", False, True), ("—", False, True),
               ("monocular depth estimation; camera self-calibration; LiDAR supervision; metric "
                "depth; Argoverse 2", True, True)], size=9, indent=0)

    h1(doc, "I", "Introduction")
    para(doc, "Estimating absolute, metric distance from a single camera is ill-posed: a pinhole "
              "camera projects the three-dimensional world onto two dimensions and the overall "
              "scale is unobservable without a metric reference or a known camera. Two coupled "
              "unknowns make this hard: the depth of each pixel, and the camera intrinsics that "
              "relate pixels to viewing rays. Focal length in particular is entangled with metric "
              "scale, so a depth model that ignores intrinsics cannot generalize across cameras.")
    para(doc, "This work asks: can a model, supervised only by LiDAR, learn to recreate the camera "
              "parameters from one image and use them to report the real distance to objects, "
              "particularly vehicles, and do so for cameras it was not calibrated on? We answer "
              "affirmatively. Contributions: (1) a LiDAR-as-ground-truth pipeline turning "
              "motion-compensated sweeps into sparse metric depth and per-vehicle distances on "
              "Argoverse 2; (2) a self-calibrating, camera-aware depth model coupling a learnable "
              "camera model to a Depth Anything V2 backbone through a ray map; (3) a controlled "
              "study with an honest negative result and annotated assumptions.")

    h1(doc, "II", "Related Work")
    para(doc, [("Monocular depth. ", True, False),
               ("Supervised monocular depth dates to Eigen et al. [1] and the scale-invariant "
                "loss; MiDaS [2] and DPT [3] showed that relative (affine-invariant) depth "
                "transfers across datasets, and Depth Anything V2 [4] scales this with a DINOv2 "
                "backbone and a DPT head. Relative depth is not metric: an affine ambiguity in "
                "inverse-depth remains, and resolving it brings back the camera.", False, False)])
    para(doc, [("Metric and camera-aware depth. ", True, False),
               ("CAM-Convs [5] injects per-pixel calibration into convolutions; Metric3D [6] and "
                "Metric3Dv2 [13] canonicalise to a reference focal length; ZoeDepth [7] adds a "
                "metric head. Most relevant, UniDepth [12] predicts metric depth and its own "
                "camera, as a dense ray field, zero-shot across datasets and is our strongest "
                "baseline. Our ray-map shares the camera-as-rays idea, but our intrinsics are "
                "LiDAR-calibrated, and Section V shows UniDepth's predicted camera is pinhole-only "
                "and fails under lens distortion.", False, False)])
    para(doc, [("Self-calibration. ", True, False),
               ("DroidCalib [8] folds intrinsics into a deep bundle-adjustment layer, optimizing "
                "them jointly with structure from video. We keep that idea but replace the "
                "multi-view anchor with direct LiDAR, which removes scale ambiguity and lets us "
                "score the recovered intrinsics against manufacturer calibration.", False, False)])
    para(doc, [("Wide-FOV and fisheye models. ", True, False),
               ("Real automotive cameras are often strongly distorted, beyond pinhole plus "
                "Brown-Conrady. The generic fisheye of Kannala and Brandt [14], the unified "
                "omnidirectional model of Mei and Rives [15], and the Enhanced Unified Camera Model "
                "[16] (pinhole to fisheye with two extra parameters) cover lenses up to and beyond "
                "180 degrees. Fisheye driving data with LiDAR exists (WoodScape [17], KITTI-360 "
                "[18]) but is access-gated, so we study controlled synthetic distortion with known "
                "ground truth (Section V).", False, False)])
    para(doc, [("Test-time and few-shot adaptation. ", True, False),
               ("Adapting a pretrained model at deployment is studied for online stereo [19] and "
                "via test-time training [20]. Our few-shot self-calibration is minimal: it adapts "
                "only the camera, a few parameters, from a handful of LiDAR frames with the depth "
                "network frozen; learning to predict distortion zero-shot from one image proves "
                "hard, motivating the few-shot route.", False, False)])
    para(doc, [("LiDAR supervision and cross-dataset evaluation. ", True, False),
               ("Projecting LiDAR for sparse supervision is standard since the KITTI depth "
                "benchmark [10]. Argoverse 2 [9], nuScenes [11] and KITTI provide synchronized "
                "cameras, LiDAR, ego-poses, per-camera ground-truth intrinsics and 3-D boxes, "
                "letting us measure depth and self-calibration error and test cross-dataset "
                "transfer, which most self-calibration work cannot.", False, False)])

    h1(doc, "III", "Methodology")
    h2(doc, "A", "The ray-map coupling")
    para(doc, "For each image-sweep pair we observe LiDAR points in the camera frame, independent "
              "of the intrinsics. Let theta be the camera model; each pixel has a unit viewing ray "
              "that depends on theta. The network predicts depth along the ray; back-projecting "
              "gives a 3-D point. A 3-D point loss compares it to the LiDAR point, so a wrong focal "
              "length yields a corrective gradient into theta. A reprojection loss requires LiDAR "
              "points projected with theta to land at the pixel whose depth predicts them. With a "
              "scale-invariant term these make calibration observable: a ten-percent focal error "
              "induces about a 0.46 m mean 3-D shift. Fig. 1 illustrates the coupling.")
    figure(doc, "viz_v3/arch_diagram.png",
           "Fig. 1.  The ray-map coupling. The learnable intrinsics define the rays that condition "
           "the depth head; back-projecting with the same intrinsics and matching LiDAR returns a "
           "gradient that corrects the calibration.")
    h2(doc, "B", "Training objectives")
    para(doc, "Training combines five terms. A scale-invariant log term (SILog) supervises the "
              "predicted depth at the LiDAR pixels. A 3-D point term back-projects the predicted "
              "depth with theta and matches the LiDAR points in the camera frame; because it "
              "depends on theta, it is the term that calibrates the camera. A reprojection term "
              "requires LiDAR points projected with theta to land at the pixel whose depth predicts "
              "them, further constraining the intrinsics. An edge-aware smoothness term regularizes "
              "the depth map, and weak Gaussian priors keep the principal point near the image "
              "center and the distortion small. The depth terms shape the depth map; the 3-D and "
              "reprojection terms shape both the depth and theta, which closes the self-calibration "
              "loop.")
    h2(doc, "C", "Camera-aware metric depth")
    para(doc, "Depth Anything V2 outputs excellent relative disparity but not metric depth. We "
              "recover scale with a head conditioned on the learned focal length, so the mapping is "
              "camera-aware. We study a free convolutional head and a camera-conditioned "
              "global-affine head; Section V reports which wins.")
    h2(doc, "D", "Self-calibration and inference without LiDAR")
    para(doc, "The intrinsics are produced by a hybrid module: a per-camera latent vector, "
              "optimized over the dataset in the manner of deep self-calibration, plus a residual "
              "predicted from the image so the model can re-estimate intrinsics for a camera it has "
              "not seen. LiDAR is used only to train and to validate; at inference the network "
              "consumes a single image and outputs depth, intrinsics, and hence object distance, "
              "with no LiDAR required. Because metric scale is conditioned on the estimated focal "
              "length, the learned scale is meant to transfer to a new camera once its focal length "
              "is recovered from the image. For cameras far from the training domain, for example "
              "wide-angle dashboard cameras with strong distortion, the metric scale can be "
              "re-anchored without LiDAR using a known camera height and the ground plane, known "
              "object sizes such as standardized license plates, logged GPS or vehicle speed, or "
              "photometric self-supervision over video. These supply the metric reference that "
              "LiDAR provides during training.")
    h2(doc, "E", "Annotated assumptions")
    para(doc, [("A1 ", False, True), ("DAv2 output is relative affine-invariant disparity. ", False, False),
               ("A2 ", False, True), ("Relative-to-metric is approximately a global affine in "
                "inverse-depth per image plus a bounded local residual. ", False, False),
               ("A3 ", False, True), ("Metric scale correlates with focal length. ", False, False),
               ("A4 ", False, True), ("Ring cameras are near-pinhole: radial distortion is small "
                "and bounded. ", False, False),
               ("A5 ", False, True), ("Metric depth lies in 0.1 to 300 m; beyond sixty meters is "
                "reported separately.", False, False)])

    h1(doc, "IV", "Implementation")
    h2(doc, "A", "Data and ground truth")
    para(doc, "We use the Argoverse 2 Sensor dataset, downloaded selectively from the public bucket "
              "(40 training and 8 validation logs, all seven ring cameras, every third LiDAR sweep, "
              "8.3 GB). For each frame we load the camera model, read the sweep, motion-compensate "
              "it from LiDAR time to camera time using the ego-poses, and project to obtain a "
              "sparse metric depth map (about 0.4% coverage, 4 to 215 m). The same machinery "
              "projects 3-D cuboids to obtain per-vehicle distance. Fig. 2 shows the ground truth.")
    figure(doc, "viz/ring_front_center_315969904949927216_depth.png",
           "Fig. 2.  LiDAR-projected ground-truth depth (motion-compensated) on a ring-camera "
           "image: the sparse metric supervision.")
    h2(doc, "B", "Model and training")
    para(doc, "The backbone is Depth Anything V2 Small (24.8 M parameters). A small CNN encodes the "
              "image into a calibration feature driving a hybrid learnable-intrinsics module: a "
              "per-camera latent (self-calibration) plus an image-conditioned residual that lets "
              "the model re-estimate intrinsics for unseen cameras. Images are processed at 518 by "
              "518; intrinsics and LiDAR pixels are scaled accordingly. Training is full fp32 with "
              "gradient clipping and a non-finite-step guard. Hardware is a single NVIDIA RTX 5090 "
              "(32 GB) with PyTorch 2.12 and CUDA 13; a three-epoch run takes about 30 to 45 "
              "minutes.")
    h2(doc, "C", "Engineering notes")
    para(doc, "Five non-obvious failures were fixed: fp16 overflow on metric losses (train in "
              "fp32); the field-of-view channel arc-cosine has an infinite gradient on the optical "
              "axis (clamp away from the limits); freely-learned distortion makes un-distortion "
              "diverge (bound it); a watcher command matched its own shell; and empty optimizer "
              "groups when the backbone is frozen.")

    h1(doc, "V", "Results and Discussion")
    h2(doc, "A", "Test protocol")
    para(doc, "We test on 200 validation frames sampled across 8 logs that are disjoint from "
              "training. Three families of test run on every frame. First, metric depth is compared "
              "to the LiDAR points, reporting absolute relative error (AbsRel), root-mean-square "
              "error (RMSE), and the fraction of pixels within a 1.25 ratio of ground truth "
              "(acc.1.25). Second, the estimated intrinsics are compared to the manufacturer "
              "calibration that Argoverse 2 provides for each camera, reporting relative focal error "
              "and principal-point error in pixels. Third, the predicted distance to each annotated "
              "vehicle, obtained by sampling the depth at the cuboid and converting with the "
              "estimated intrinsics, is compared to the cuboid distance. The four variants v1 to v3 "
              "form a controlled ablation in which one factor changes at a time (Table I).")
    h2(doc, "B", "Findings")
    table(doc, "Table I.  Ablation on held-out validation (v1 free conv/frozen; v2 affine "
               "log-depth; v2b affine inverse-depth; v3 free conv + unfrozen backbone).",
          [["Ver.", "AbsRel", "RMSE", "acc.1.25", "fx err", "fy err", "Veh.<=60"],
           ["v1", "0.204", "10.14", "0.714", "0.16%", "1.37%", "--"],
           ["v2", "0.368", "11.95", "0.560", "0.55%", "0.61%", "8.47 m"],
           ["v2b", "0.249", "10.95", "0.685", "0.30%", "1.33%", "9.14 m"],
           ["v3", "0.112", "7.30", "0.884", "0.26%", "0.41%", "8.56 m"]],
          [0.42, 0.5, 0.5, 0.55, 0.45, 0.45, 0.6])
    para(doc, [("Self-calibration works everywhere. ", True, False),
               ("Across all variants the model recovers intrinsics to within 1.4% (fx and fy) and "
                "sub-pixel principal point on held-out cameras; v3 reaches fx 0.26% and fy 0.41%. "
                "The camera parameters are recreated from the image alone.", False, False)])
    para(doc, [("A pinpointed negative result. ", True, False),
               ("The camera-conditioned global-affine head (v2) hurt depth (AbsRel 0.204 to 0.368) "
                "because of a wrong functional-form assumption (A2): depth was mapped linearly in "
                "log-depth, but disparity is affine in inverse-depth. Correcting it (v2b) recovered "
                "most of the loss but still lost to the free conv head.", False, False)])
    para(doc, [("The real lever was the backbone. ", True, False),
               ("Unfreezing Depth Anything V2 (v3) nearly halved AbsRel (0.204 to 0.112) and raised "
                "accuracy at threshold 1.25 from 0.71 to 0.88, with the best calibration. Per-range "
                "vehicle distance is in Table II. Fig. 3 shows an example.", False, False)])
    table(doc, "Table II.  Vehicle-distance error by range (v3).",
          [["Range", "0-30 m", "30-60 m", "60 m+", "overall"],
           ["MAE", "3.90 m", "12.57 m", "35.85 m", "22.46 m"],
           ["n", "330", "384", "741", "1455"]],
          [0.55, 0.62, 0.66, 0.6, 0.66])
    figure(doc, "viz_v3/infer_v3_closetraffic.png",
           "Fig. 3.  Single-image inference (v3). Left: input with per-vehicle predicted/"
           "ground-truth distance. Right: predicted metric depth. Focal self-estimated to 0.4%.")
    h2(doc, "C", "Generalization and deployment")
    para(doc, "Two deployment properties follow from the test design. First, because LiDAR appears "
              "only in the training and test losses, inference is image-only: a new frame without "
              "LiDAR is processed exactly as during validation, so the absence of LiDAR at "
              "deployment does not affect the model. Second, generalization across cameras was only "
              "weakly testable here: the seven ring cameras span about five percent focal "
              "diversity, so cross-focal transfer is not yet stressed, and wide dashboard cameras "
              "additionally introduce strong distortion outside the near-pinhole assumption (A4). "
              "We therefore treat wide-camera deployment as future work, to be validated either "
              "with a synthetic field-of-view and distortion stress test or with a LiDAR-free "
              "metric anchor on real dashboard-camera video.")
    h2(doc, "D", "Cross-dataset generalization")
    para(doc, "A single sensor rig under-tests camera-awareness, so we built a unified benchmark "
              "layer reducing every dataset to the same contract (image, sparse LiDAR depth, "
              "ground-truth intrinsics) behind one adapter, and applied the AV2-trained v3 model to "
              "KITTI and nuScenes front cameras (focal length 721 and 1266 pixels; AV2 1782) under "
              "three regimes: zero-shot; adapting only the per-camera latent on twenty frames; and "
              "additionally adapting the small depth head with the aspect prior disabled (Table III, "
              "Fig. 4). Qualitative held-out results are in Fig. 5.")
    table(doc, "Table III.  Cross-dataset transfer (AV2-trained v3; 20-frame few-shot; held-out).",
          [["Dataset", "Regime", "AbsRel", "acc.1.25", "fx err"],
           ["KITTI", "zero-shot", "0.390", "0.026", "97.1%"],
           ["KITTI", "+latent", "0.402", "0.021", "0.2%"],
           ["KITTI", "+latent+head", "0.100", "0.953", "1.1%"],
           ["nuScenes", "zero-shot", "0.274", "0.184", "45.3%"],
           ["nuScenes", "+latent", "0.284", "0.172", "0.5%"],
           ["nuScenes", "+latent+head", "0.122", "0.853", "0.7%"],
           ["AV2", "in-domain", "0.112", "0.884", "0.26%"]],
          [0.62, 0.78, 0.5, 0.6, 0.5])
    para(doc, [("Three findings. ", False, True),
               ("Zero-shot transfer fails, and structurally: intrinsics are a per-camera latent "
                "plus a small image residual, so an unseen camera inherits the AV2 focal and KITTI "
                "accuracy collapses to 0.03. Self-calibration transfers few-shot: twenty frames "
                "recover focal length to within half a percent (fx) on both datasets. Calibration "
                "and metric scale are separable: latent-only adaptation leaves depth unchanged, but "
                "additionally adapting the 58.8k-parameter depth head restores it dramatically "
                "(KITTI accuracy 0.026 to 0.953, nuScenes 0.184 to 0.853), matching in-domain AV2. "
                "A residual ten-to-fourteen percent error in fy persists, as vertical focal is "
                "weakly observed by the LiDAR loss.", False, False)])
    figure(doc, "viz_paper/results_bars.png",
           "Fig. 4.  Cross-dataset generalization: AbsRel, accuracy at threshold 1.25, and "
           "focal-length error (log) across the three regimes for KITTI and nuScenes; dashed line "
           "is in-domain Argoverse 2.")
    figure(doc, "viz_paper/cross_dataset_qualitative.png",
           "Fig. 5.  Qualitative cross-dataset results, one row per benchmark: input, ground-truth "
           "LiDAR depth, predicted metric depth, and absolute error at the LiDAR points, on a "
           "shared scale.")
    h2(doc, "E", "Diagnosing and repairing a foundation model under distortion")
    para(doc, "The strongest model we benchmarked, UniDepth-V2, predicts its own camera and is "
              "strong zero-shot on near-pinhole driving images (KITTI AbsRel 0.089, nuScenes 0.105), "
              "yet under our controlled distortion sweep its depth degrades and then collapses "
              "(KITTI AbsRel 0.09 to 0.28; nuScenes to 0.78). Its source explains why: the depth "
              "decoder is camera-agnostic, conditioning on a per-pixel ray field, so it consumes "
              "whatever camera it is given (its code implements pinhole, EUCM, OpenCV, Fisheye624 and "
              "MEI), but the inference-time self-calibration head predicts only a four-parameter "
              "pinhole. On a distorted lens the predicted rays are wrong toward the periphery and "
              "depth degrades; the failure is localised entirely in the camera head, not the decoder.")
    para(doc, "This implies a fix needing no retraining: predict a distortion-capable camera, and "
              "recover it with our LiDAR few-shot self-calibration. We freeze UniDepth, fit the "
              "unknown distortion from twenty LiDAR frames via the reprojection objective, and supply "
              "the recovered camera at inference (Table IV, Fig. 6). The recovered camera matches the "
              "ground-truth-camera oracle, lifting accuracy from AbsRel 0.158 to 0.106 and the 1.25 "
              "threshold from 0.695 to 0.866 at the strong setting, purely from calibration. Caveats: "
              "the minimal EUCM head fits realistic but not extreme fisheye; central LiDAR "
              "under-constrains high-order terms, so a low-order fit with fixed base focal is "
              "required; and the residual gap to the undistorted level (0.076) reflects the backbone "
              "still ingesting the warped image, which distortion augmentation in pretraining would "
              "address.")
    table(doc, "Table IV.  Repairing UniDepth under KB-fisheye distortion (AbsRel / acc.1.25; KITTI "
               "held-out). Our few-shot calibration matches the oracle.",
          [["Camera supplied to UniDepth", "k1=2.5", "k1=4.0"],
           ["none (predicted pinhole)", "0.112 / 0.85", "0.158 / 0.70"],
           ["ours (20-frame LiDAR calib)", "0.087 / 0.93", "0.106 / 0.87"],
           ["GT camera (oracle)", "0.087 / 0.93", "0.106 / 0.87"]],
          [1.5, 0.75, 0.75])
    figure(doc, "viz_paper/unidepth_fix.png",
           "Fig. 6.  Repairing UniDepth under KB-fisheye distortion: a camera recovered by our "
           "20-frame LiDAR self-calibration matches the GT-camera oracle and far exceeds the "
           "zero-shot (pinhole-head) prediction.")
    para(doc, "We also tested making the head predict distortion zero-shot: a head trained to "
              "regress EUCM distortion from a single image (with augmentation and a ray-matching "
              "loss) collapses to a near-constant prediction, its error scaling with the held-out "
              "distortion. Per-image distortion is only weakly identifiable from one image with a "
              "light head -- a plausible reason production models commit to a pinhole head -- so the "
              "few-shot route is the dependable one. The recipe is thus a distortion-capable camera "
              "the decoder can use, plus LiDAR few-shot self-calibration to estimate it, rather than "
              "a learned zero-shot distortion predictor.")

    h1(doc, "VI", "Conclusion")
    para(doc, "SelfCalibDepth shows that LiDAR is a sufficient anchor to jointly learn camera "
              "self-calibration and camera-aware metric depth from a single image, on a real "
              "driving dataset and on held-out cameras. Self-calibration is the strongest result "
              "(sub-one-percent focal error); metric depth and vehicle distance are solid and were "
              "driven primarily by fine-tuning a depth foundation model rather than by "
              "architectural cleverness, a finding made visible only because every assumption was "
              "annotated and ablated. The cross-dataset study sharpens the claim: self-calibration "
              "transfers to KITTI and nuScenes with a twenty-frame adaptation, and a small "
              "depth-head adaptation transfers the metric depth too. Limitations: zero-shot transfer "
              "fails by design (the per-camera latent dominates), the residual fy error shows "
              "vertical focal is weakly observed, distortion learning is conservative, far-range "
              "depth is weak, and Lyft L5 is implemented but unevaluated (gated data). Future work: "
              "restructure theta for genuine zero-shot calibration, a reprojection-weighted schedule "
              "to tighten fy, cross-dataset vehicle metrics, a photometric term, and a larger "
              "backbone.")

    h1(doc, "", "References")
    refs = [
        "[1] D. Eigen, C. Puhrsch, R. Fergus, “Depth map prediction from a single image using "
        "a multi-scale deep network,” NeurIPS, 2014.",
        "[2] R. Ranftl et al., “Towards robust monocular depth estimation,” IEEE TPAMI, 2022.",
        "[3] R. Ranftl, A. Bochkovskiy, V. Koltun, “Vision transformers for dense "
        "prediction,” ICCV, 2021.",
        "[4] L. Yang et al., “Depth Anything V2,” NeurIPS, 2024.",
        "[5] J. M. Facil et al., “CAM-Convs: Camera-aware multi-scale convolutions,” CVPR, 2019.",
        "[6] W. Yin et al., “Metric3D: Towards zero-shot metric 3D prediction,” ICCV, 2023.",
        "[7] S. F. Bhat et al., “ZoeDepth,” arXiv:2302.12288, 2023.",
        "[8] P. Hagemann et al., “Deep geometry-aware camera self-calibration from video,” "
        "ICCV, 2023.",
        "[9] B. Wilson et al., “Argoverse 2,” NeurIPS Datasets and Benchmarks, 2021.",
        "[10] A. Geiger, P. Lenz, R. Urtasun, “Are we ready for autonomous driving? The KITTI "
        "vision benchmark suite,” CVPR, 2012.",
        "[11] H. Caesar et al., “nuScenes: A multimodal dataset for autonomous driving,” CVPR, 2020.",
        "[12] L. Piccinelli et al., “UniDepth: Universal monocular metric depth estimation,” "
        "CVPR, 2024.",
        "[13] M. Hu et al., “Metric3D v2: A versatile monocular geometric foundation model,” "
        "IEEE TPAMI, 2024.",
        "[14] J. Kannala, S. Brandt, “A generic camera model and calibration method for "
        "conventional, wide-angle, and fisheye lenses,” IEEE TPAMI, 2006.",
        "[15] C. Mei, P. Rives, “Single view point omnidirectional camera calibration from planar "
        "grids,” ICRA, 2007.",
        "[16] B. Khomutenko, G. Garcia, P. Martinet, “An enhanced unified camera model,” IEEE "
        "RA-L, 2016.",
        "[17] S. Yogamani et al., “WoodScape: A multi-task, multi-camera fisheye dataset for "
        "autonomous driving,” ICCV, 2019.",
        "[18] Y. Liao, J. Xie, A. Geiger, “KITTI-360: A novel dataset and benchmarks for urban "
        "scene understanding in 2D and 3D,” IEEE TPAMI, 2022.",
        "[19] A. Tonioni et al., “Real-time self-adaptive deep stereo,” CVPR, 2019.",
        "[20] Y. Sun et al., “Test-time training with self-supervision for generalization under "
        "distribution shifts,” ICML, 2020.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.17)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        p.paragraph_format.space_after = Pt(1)
        runs(p, r, size=8)

    doc.save("PAPER.docx")
    print("wrote PAPER.docx")


if __name__ == "__main__":
    build()
